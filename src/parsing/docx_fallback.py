"""DOCX fallback parser using python-docx (placeholder for future use)."""
from pathlib import Path
from typing import Any


def parse_with_docx(file_path: Path) -> dict[str, Any]:
    """Parse DOCX using python-docx. Used only if Docling unavailable."""
    from docx import Document

    doc = Document(str(file_path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    tables = _extract_tables(doc)

    markdown_parts = list(paragraphs)
    for table in tables:
        grid = table["data"]["grid"]
        if grid:
            markdown_parts.append(_grid_to_markdown(grid))
    markdown = "\n\n".join(markdown_parts)

    return {
        "markdown": markdown,
        "doc_dict": {"paragraphs": paragraphs, "tables": tables},
        "sections": [],
        "tables": tables,
        "figures": [],
        "page_count": 1,
        "parser": "python-docx",
        "source_format": ".docx",
    }


def _extract_tables(doc) -> list[dict[str, Any]]:
    tables = []
    for i, table in enumerate(doc.tables):
        grid = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if not any(any(cell for cell in row) for row in grid):
            continue
        tables.append({
            "table_id": f"table_001_{i:02d}",
            "page": 1,
            "data": {
                "grid": grid,
                "num_rows": len(grid),
                "num_cols": max((len(row) for row in grid), default=0),
            },
        })
    return tables


def _grid_to_markdown(grid: list[list[str]]) -> str:
    if not grid:
        return ""
    width = max(len(row) for row in grid)
    rows = [row + [""] * (width - len(row)) for row in grid]
    lines = [
        "| " + " | ".join(rows[0]) + " |",
        "| " + " | ".join("---" for _ in range(width)) + " |",
    ]
    for row in rows[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)
